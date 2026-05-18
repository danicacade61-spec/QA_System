"""
生成中期报告Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_borders(table):
    """为表格添加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_paragraph_with_format(doc, text, style_name=None, font_size=None, bold=False, 
                               alignment=None, first_line_indent=None, space_before=None,
                               space_after=None, font_name=None, color=None):
    """添加格式化的段落"""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    
    run = p.add_run(text)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = bold
    if font_name:
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    
    return p

def create_cover_page(doc):
    """创建封面"""
    # 空行
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    
    # 学校名称
    add_paragraph_with_format(doc, "西南交通大学", font_size=28, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_after=4)
    
    # 学院名称
    add_paragraph_with_format(doc, "计算机与人工智能学院", font_size=20, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_after=40)
    
    # 分隔线和报告类型
    add_paragraph_with_format(doc, "课程设计中期报告", font_size=26, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=30, space_after=20)
    
    # 空行
    doc.add_paragraph()
    
    # 设计题目
    add_paragraph_with_format(doc, "设计题目", font_size=16, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_after=8)
    
    add_paragraph_with_format(doc, "基于大语言模型的智能问答方法研究与实现", font_size=20, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_after=40)
    
    # 信息表
    info_items = [
        ("年    级：", "2025级"),
        ("学    院：", "计算机与人工智能学院"),
        ("专    业：", "计算机科学与技术"),
        ("学号姓名：", "2025212413    毛树林"),
        ("课程名称：", "自然语言处理"),
        ("指导教师：", "贾真"),
    ]
    
    info_table = doc.add_table(rows=len(info_items), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(info_table)
    
    for i, (label, value) in enumerate(info_items):
        cell_label = info_table.cell(i, 0)
        cell_value = info_table.cell(i, 1)
        
        # 设置列宽
        cell_label.width = Cm(4)
        cell_value.width = Cm(8)
        
        # 标签单元格
        p_label = cell_label.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_label = p_label.add_run(label)
        run_label.font.size = Pt(14)
        run_label.font.name = "宋体"
        
        # 值单元格
        p_value = cell_value.paragraphs[0]
        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_value = p_value.add_run(value)
        run_value.font.size = Pt(14)
        run_value.font.name = "宋体"
        run_value.font.bold = True
    
    # 日期
    doc.add_paragraph()
    add_paragraph_with_format(doc, "二零二六年五月", font_size=18, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=30)
    
    # 分页
    doc.add_page_break()


def create_toc_page(doc):
    """创建目录页"""
    add_paragraph_with_format(doc, "目  录", font_size=22, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_after=20)
    
    toc_items = [
        ("摘要", 1),
        ("一、项目概述与开题回顾", 1),
        ("  1.1 项目目标", 2),
        ("  1.2 开题计划回顾", 2),
        ("二、已完成工作", 1),
        ("  2.1 系统架构设计", 2),
        ("  2.2 核心模块实现", 2),
        ("  2.3 前后端实现", 2),
        ("  2.4 评估模块实现", 2),
        ("  2.5 代码质量优化", 2),
        ("三、进度完成情况", 1),
        ("  3.1 实际进度对照", 2),
        ("  3.2 已实现功能清单", 2),
        ("四、当前问题与解决方案", 1),
        ("五、后续工作计划", 1),
        ("  5.1 下一阶段安排", 2),
        ("  5.2 预期成果", 2),
        ("六、参考文献", 1),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.name = "宋体"
        if level == 1:
            run.font.bold = True
        p.paragraph_format.first_line_indent = Pt(0)


def create_body_content(doc):
    """创建正文内容"""
    
    # ===== 摘要 =====
    add_paragraph_with_format(doc, "摘要", font_size=22, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=10, space_after=10)
    
    abstract_text = (
        "本中期报告总结了《基于大语言模型的智能问答方法研究与实现》课程设计的阶段性进展。"
        "截至目前，项目已按计划完成文献调研与原理分析、系统架构设计、核心模块开发等关键任务，"
        "成功构建了一个基于检索增强生成（RAG）架构的智能问答系统原型。系统实现了涵盖文档加载与分块、"
        "文本向量化与FAISS索引存储、向量相似度检索、大语言模型调用与回答生成、RESTful API服务以及"
        "Streamlit交互界面等的完整功能链条，支持智谱GLM API、OpenAI API和本地Qwen2模型三种LLM后端方案。"
        "系统采用模块化分层架构（core/api/services/ui/utils/models），实现了策略模式、工厂模式和延迟加载等设计模式。"
        "评估模块已开发完成，支持精确匹配（Exact Match）、Token F1、ROUGE-L等多种自动评估指标。"
        "代码质量方面，已完成loguru日志系统引入、jieba中文分词优化、Pydantic模型重构、ZhipuClient配置统一等优化工作。"
        "后续工作将聚焦于知识库建设、系统性测试评估与课程设计报告撰写。"
    )
    add_paragraph_with_format(doc, abstract_text, font_size=14, font_name="宋体",
                              first_line_indent=0.74, space_after=6)
    
    add_paragraph_with_format(doc, 
        "关键词：大语言模型；智能问答；检索增强生成；RAG；自然语言处理",
        font_size=14, font_name="宋体", first_line_indent=0, space_after=16)
    
    add_paragraph_with_format(doc, "Abstract", font_size=22, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="Times New Roman",
                              space_before=10, space_after=10)
    
    abstract_en = (
        "This interim report summarizes the progress of the course project "
        "\"Research and Implementation of Intelligent Question Answering Methods Based on Large Language Models.\" "
        "To date, the project has completed key tasks including literature review and theoretical analysis, "
        "system architecture design, and core module development, successfully building a prototype intelligent "
        "QA system based on the Retrieval-Augmented Generation (RAG) architecture. The system implements the "
        "complete functional pipeline: document loading and chunking, text vectorization with FAISS indexing, "
        "vector similarity retrieval, LLM invocation and answer generation, RESTful API services, and a "
        "Streamlit interactive interface. It supports three LLM backend solutions: Zhipu GLM API, OpenAI API, "
        "and local Qwen2 models. The system adopts a modular layered architecture with design patterns including "
        "Strategy, Factory Method, and Lazy Loading. The evaluation module supports automatic metrics such as "
        "Exact Match, Token F1, and ROUGE-L. Code quality improvements include loguru logging, jieba-based "
        "Chinese tokenization, Pydantic model refactoring, and ZhipuClient configuration unification. "
        "Future work will focus on knowledge base construction, systematic testing and evaluation, and "
        "report writing."
    )
    add_paragraph_with_format(doc, abstract_en, font_size=12, font_name="Times New Roman",
                              first_line_indent=0, space_after=6)
    
    add_paragraph_with_format(doc, 
        "Keywords: Large Language Model; Question Answering; Retrieval-Augmented Generation; RAG; NLP",
        font_size=12, font_name="Times New Roman", first_line_indent=0, space_after=16)
    
    doc.add_page_break()
    
    # ===== 一、项目概述与开题回顾 =====
    add_paragraph_with_format(doc, "一、项目概述与开题回顾", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    add_paragraph_with_format(doc, "1.1 项目目标", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)

    add_paragraph_with_format(doc,
        "本课题来源于《自然语言处理》课程设计，选题为“基于大语言模型的智能问答方法研究与实现”。"
        "研究目标是研究基于大语言模型的智能问答方法的实现原理，并实现一个完整的智能问答系统原型。"
        "围绕这一目标，本课题聚焦以下两个核心研究问题：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)

    problems = [
        "问题一：大语言模型实现智能问答的原理是什么？具体包括Transformer架构如何处理自然语言输入、预训练如何赋予模型问答能力、自回归生成如何逐步产出回答。",
        "问题二：如何实现一个基于大语言模型的智能问答系统？具体包括模型部署、后端推理服务搭建、前端交互界面实现、系统集成与功能测试。"
    ]
    for i, problem in enumerate(problems):
        add_paragraph_with_format(doc, f"（{i+1}）{problem}", font_size=14,
                                  font_name="宋体", first_line_indent=0.74, space_after=2)

    add_paragraph_with_format(doc,
        "围绕上述研究问题，本课题的研究内容从实现原理和系统实现两个层面展开：一是系统梳理大语言模型"
        "实现智能问答的技术原理，包括Transformer架构、预训练机制和自回归生成过程；二是选定开源大语言模型"
        "作为底座，参考RAG等智能问答系统架构，设计并实现一个包含模型部署、后端推理服务和前端交互界面的"
        "完整智能问答系统原型。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # 表1 开题计划
    add_paragraph_with_format(doc, "表1 开题计划进度表", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    plan_table = doc.add_table(rows=5, cols=4)
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(plan_table)
    
    headers = ["阶段", "时间安排", "主要任务", "预期成果"]
    for i, h in enumerate(headers):
        cell = plan_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    plan_data = [
        ["第一阶段", "第5~6周", "文献调研，理解实现原理，选定开源模型，确定技术方案", "完成文献综述，确定技术方案"],
        ["第二阶段", "第7~10周", "智能问答系统设计与开发", "系统原型可运行"],
        ["第三阶段", "第11~13周", "系统测试、效果验证与优化", "系统功能完善，通过测试"],
        ["第四阶段", "第14~17周", "课程设计报告撰写与修改", "完成课程设计报告终稿"],
    ]
    for row_idx, row_data in enumerate(plan_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = plan_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            run.font.name = "宋体"
    
    doc.add_page_break()
    
    # ===== 二、已完成工作 =====
    add_paragraph_with_format(doc, "二、已完成工作", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    add_paragraph_with_format(doc, "2.1 系统架构设计", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc, 
        "系统采用经典的检索增强生成（RAG）架构，整体流程为：用户提问 → 文档检索（向量相似度搜索）→ "
        "上下文构建 → 大语言模型生成回答。系统基于模块化分层设计，主要包含以下层级：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    layers = [
        "数据层：文档加载与解析（TXT/PDF/DOCX/MD）、文本分块处理、向量化与索引存储",
        "检索层：基于FAISS的向量相似度检索、相关性排序与过滤",
        "生成层：多LLM后端支持（OpenAI/智谱/本地模型）、上下文增强提示词构建、流式生成",
        "服务层：FastAPI RESTful API、知识库管理服务、对话历史管理",
        "展示层：Streamlit交互式Web界面（含对话、知识库管理、检索预览三标签页）",
    ]
    for layer in layers:
        add_paragraph_with_format(doc, layer, font_size=14, font_name="宋体",
                                  first_line_indent=0.74, space_after=2)
    
    add_paragraph_with_format(doc, "2.2 核心模块实现", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    # 文档处理模块
    add_paragraph_with_format(doc, "2.2.1 文档处理模块（DocumentProcessor）", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "实现了对TXT、PDF、DOCX、MD四种格式文档的自动加载与解析。对于TXT和MD文件直接读取文本内容；"
        "PDF使用pypdf库逐页提取文字；DOCX使用python-docx库提取段落文本。分块策略采用段落级与句子级相结合"
        "的混合分割方式：先按空行划分段落，再对超过块大小（默认256字符）的长段落按句号、问号、感叹号等句子边界"
        "进一步拆分，同时支持块间重叠（默认32字符）以保持上下文连贯性。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # 嵌入向量管理模块
    add_paragraph_with_format(doc, "2.2.2 嵌入向量管理模块（EmbeddingManager）", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "采用BAAI/bge-small-zh-v1.5作为文本嵌入模型，输出512维向量，在中文语义表示任务上表现优良。"
        "通过sentence-transformers库加载模型，支持批量编码以提高效率。编码后的向量自动进行L2归一化处理，"
        "使得使用内积（Inner Product）进行相似度检索等价于余弦相似度计算。模块采用延迟加载策略，"
        "在首次使用时才加载模型，避免系统启动时不必要的资源消耗。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # 向量存储模块
    add_paragraph_with_format(doc, "2.2.3 向量存储模块（VectorStore）", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "基于FAISS库实现向量数据库，使用IndexIDMap包装IndexFlatIP索引结构，支持向量添加、相似度检索、"
        "删除和清空等操作。检索时设置相似度阈值（默认0.3），仅返回超过阈值的文档块以保证检索质量。"
        "索引数据持久化到磁盘，包括FAISS索引文件、ID映射文件和元数据文件，系统重启时自动加载已有索引。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # LLM客户端模块
    add_paragraph_with_format(doc, "2.2.4 大语言模型客户端模块（LLMClient）", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "采用策略模式和工厂方法设计模式，定义统一的BaseLLMClient抽象基类，提供了chat()和chat_stream()"
        "两个核心接口。实现了三个具体客户端：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=2)
    
    clients = [
        "OpenAIClient：通过HTTPX库调用兼容OpenAI格式的API，支持流式与非流式两种调用方式",
        "ZhipuClient：调用智谱ChatGLM API（当前配置使用glm-5.1模型），接口格式与OpenAI兼容",
        "LocalModelClient：使用Transformers库加载开源模型（默认Qwen2-1.5B-Instruct），支持CPU/GPU推理",
    ]
    for c in clients:
        add_paragraph_with_format(doc, f"（3）{c}", font_size=14, font_name="宋体",
                                  first_line_indent=0.74, space_after=2)
    
    add_paragraph_with_format(doc, 
        "客户端通过create_llm_client()工厂函数根据配置文件自动创建对应的实例。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # RAG引擎模块
    add_paragraph_with_format(doc, "2.2.5 RAG引擎模块（RAGEngine）", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "RAG引擎是整个系统的核心调度模块，实现了完整的检索增强生成流程：接收用户问题后，调用VectorStore"
        "进行向量相似度检索获取最相关的top-k个文档块，将检索结果拼接为上下文构建系统提示词，结合对话历史"
        "调用LLM客户端生成回答，最后更新对话历史并返回结构化结果。引擎同时实现了流式查询接口，支持逐词返回生成结果。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    add_paragraph_with_format(doc, "2.3 前后端实现", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc, "2.3.1 FastAPI后端服务", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "基于FastAPI框架构建RESTful API服务器，部署于8000端口。实现了问答查询（含流式SSE响应）、"
        "知识库管理（统计、上传文档、删除文档、清空知识库）、对话历史管理、系统配置查询等API端点。"
        "API配备完整的CORS跨域支持和Pydantic数据模型验证，自动生成Swagger交互式文档。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # API端点表
    add_paragraph_with_format(doc, "表2 API端点一览", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    api_table = doc.add_table(rows=10, cols=3)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(api_table)
    
    api_headers = ["端点", "方法", "说明"]
    for i, h in enumerate(api_headers):
        cell = api_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    api_data = [
        ["/", "GET", "系统状态与版本信息"],
        ["/api/query", "POST", "问答查询（非流式）"],
        ["/api/query/stream", "POST", "流式问答（SSE）"],
        ["/api/chat/clear", "POST", "清空对话历史"],
        ["/api/knowledge-base/stats", "GET", "知识库统计信息"],
        ["/api/knowledge-base/upload", "POST", "上传文档"],
        ["/api/knowledge-base/document/{name}", "DELETE", "删除指定文档"],
        ["/api/knowledge-base/clear", "POST", "清空知识库"],
        ["/api/config", "GET", "获取系统配置"],
    ]
    for row_idx, row_data in enumerate(api_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = api_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "宋体"
    
    add_paragraph_with_format(doc, "2.3.2 Streamlit前端界面", font_size=14, bold=True,
                              font_name="宋体", space_before=8, space_after=4)
    
    add_paragraph_with_format(doc, 
        "采用Streamlit框架构建了交互式Web界面，部署于8501端口。界面包含三个功能标签页：对话标签页支持"
        "多轮对话和RAG检索开关，采用气泡式聊天样式；知识库管理标签页支持单文件上传和批量目录导入；"
        "检索预览标签页可直观查看向量检索结果。侧边栏集成了系统状态面板、检索参数设置和对话控制功能。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    add_paragraph_with_format(doc, "2.4 评估模块实现", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc,
        "评估模块（QAEvaluator）支持加载多种格式的测试数据（JSON/JSONL/TXT），对问答结果进行自动评估。"
        "实现了精确匹配（Exact Match）、Token F1分数、ROUGE-L（基于最长公共子序列）等自动评估指标，"
        "同时记录响应时间和检索覆盖率等性能指标。评估结果自动保存为JSON和TXT格式的报告，支持批量评估和汇总统计。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)

    add_paragraph_with_format(doc, "2.5 代码质量优化", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)

    add_paragraph_with_format(doc,
        "在系统核心功能开发完成后，对代码进行了全面的质量审查与优化，主要改进包括："
        "（1）引入loguru日志框架替换原有print()输出，按照日志级别（INFO/WARNING/ERROR）统一管理所有核心模块的运行时信息；"
        "（2）优化嵌入管理器的中文分词降级方案，优先使用jieba进行中文分词以提高分词质量；"
        "（3）重构API层数据模型，将Pydantic模型从server.py中独立到app/models/schemas.py，提升代码组织性；"
        "（4）统一ZhipuClient的API地址配置，改为从Settings中读取BASE_URL，与OpenAIClient行为保持一致；"
        "（5）修正LocalModelClient的对话模板为Qwen2标准格式（<|im_start|>），提升本地模型兼容性；"
        "（6）实现PII过滤工具（filter_pii.py），支持扫描和替换代码中的个人敏感信息；"
        "（7）补充.env.example环境变量配置示例，降低项目上手门槛。",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)

    doc.add_page_break()
    
    # ===== 三、进度完成情况 =====
    add_paragraph_with_format(doc, "三、进度完成情况", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    add_paragraph_with_format(doc, "3.1 实际进度对照", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    # 表3 进度完成情况
    add_paragraph_with_format(doc, "表3 进度完成情况对照表", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    progress_table = doc.add_table(rows=4, cols=4)
    progress_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(progress_table)
    
    prog_headers = ["阶段", "计划任务", "完成情况", "完成度"]
    for i, h in enumerate(prog_headers):
        cell = progress_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    prog_data = [
        ["第一阶段", "文献调研、原理理解、技术方案确定", 
         "已完成文献综述研究，理解了Transformer、预训练、自回归生成等技术原理，确定了RAG架构作为技术方案",
         "100%"],
        ["第二阶段", "系统设计与开发",
         "已完成全部核心模块编码：文档处理、嵌入管理、向量存储、LLM客户端、RAG引擎、API服务、前端界面、评估模块",
         "约95%"],
        ["第三阶段", "系统测试、效果验证与优化",
         "评估模块已开发完成，测试数据集已创建（10个样本），系统性测试评估部分进行中",
         "约30%"],
    ]
    for row_idx, row_data in enumerate(prog_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = progress_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            run.font.name = "宋体"
            if col_idx == 3:
                if row_idx < 2:
                    run.font.color.rgb = RGBColor(16, 185, 129)
                else:
                    run.font.color.rgb = RGBColor(245, 158, 11)
    
    add_paragraph_with_format(doc, "3.2 已实现功能清单", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc, 
        "以下为项目目前已实现的主要功能点：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # 表4 功能实现清单
    add_paragraph_with_format(doc, "表4 功能实现清单", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    func_table = doc.add_table(rows=11, cols=3)
    func_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(func_table)
    
    func_headers = ["功能模块", "功能点", "状态"]
    for i, h in enumerate(func_headers):
        cell = func_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    func_data = [
        ["文档处理", "TXT/PDF/DOCX/MD文档加载与智能分块", "已完成"],
        ["嵌入与向量", "BAAI/bge-small-zh-v1.5嵌入 + FAISS索引", "已完成"],
        ["LLM客户端", "OpenAI/智谱API/本地模型三种后端", "已完成"],
        ["RAG引擎", "检索增强生成 + 流式输出 + 多轮对话", "已完成"],
        ["API服务", "FastAPI RESTful API + SSE流式响应", "已完成"],
        ["前端界面", "Streamlit对话/知识库管理/检索预览", "已完成"],
        ["评估模块", "Exact Match / Token F1 / ROUGE-L", "已完成"],
        ["知识库管理", "上传/删除/清空/统计", "已完成"],
        ["系统配置", "多提供者切换/参数可配置/延迟加载", "已完成"],
        ["代码质量优化", "loguru日志/jieba分词/模型重构/配置统一", "已完成"],
    ]
    for row_idx, row_data in enumerate(func_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = func_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            run.font.name = "宋体"
    
    doc.add_page_break()
    
    # ===== 四、当前问题与解决方案 =====
    add_paragraph_with_format(doc, "四、当前问题与解决方案", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    # 表5 问题
    add_paragraph_with_format(doc, "表5 当前问题与应对策略", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    issue_table = doc.add_table(rows=6, cols=4)
    issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(issue_table)
    
    issue_headers = ["序号", "问题描述", "影响", "解决方案"]
    for i, h in enumerate(issue_headers):
        cell = issue_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    issue_data = [
        ["1", "知识库目前为空，尚未加载文档", "RAG检索功能无法完整验证", "已开发完成知识库构建功能，后续上传课程资料"],
        ["2", "测试数据规模偏小（仅10个样本）", "评估结果统计意义有限", "计划扩展测试数据集至100-200条"],
        ["3", "本地模型推理速度较慢", "影响用户体验", "以API调用为主要方案，本地模型作为备选"],
        ["4", "尚未进行系统性人工评估", "回答质量缺乏人工维度验证", "将从准确性、相关性、流畅性三维度评分"],
        ["5", "文档分块策略有待优化", "影响检索准确率", "计划对比不同分块大小和重叠策略的效果"],
    ]
    for row_idx, row_data in enumerate(issue_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = issue_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            run.font.name = "宋体"
    
    # ===== 五、后续工作计划 =====
    add_paragraph_with_format(doc, "五、后续工作计划", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    add_paragraph_with_format(doc, "5.1 下一阶段安排", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc, 
        "结合当前进度，第三、四阶段的具体工作计划如下：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    # 表6 后续工作
    add_paragraph_with_format(doc, "表6 后续工作安排", font_size=12, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体",
                              space_before=6, space_after=4)
    
    plan2_table = doc.add_table(rows=6, cols=3)
    plan2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(plan2_table)
    
    plan2_headers = ["周次", "工作任务", "预期成果"]
    for i, h in enumerate(plan2_headers):
        cell = plan2_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = "宋体"
        set_cell_shading(cell, "D9E2F3")
    
    plan2_data = [
        ["第11周", "构建知识库、扩展测试数据", "知识库具备初步规模，测试数据集扩充"],
        ["第12周", "系统性测试评估、参数调优", "获得完整评估报告，完成参数优化"],
        ["第13周", "人工评估、系统优化与Bug修复", "系统功能完善，通过全面测试"],
        ["第14~15周", "撰写课程设计报告初稿", "完成设计报告初稿"],
        ["第16~17周", "报告修改完善、准备答辩材料", "完成设计报告终稿"],
    ]
    for row_idx, row_data in enumerate(plan2_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = plan2_table.cell(row_idx + 1, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            run.font.name = "宋体"
    
    add_paragraph_with_format(doc, "5.2 预期成果", font_size=16, bold=True,
                              font_name="宋体", space_before=10, space_after=6)
    
    add_paragraph_with_format(doc, 
        "项目最终预期达成的成果包括：",
        font_size=14, font_name="宋体", first_line_indent=0.74, space_after=4)
    
    outcomes = [
        "一个可运行的智能问答系统原型：支持通过Web界面和API进行智能问答交互，具备知识库管理与检索增强生成能力",
        "系统性测试评估报告：包含自动指标（准确率、F1、ROUGE-L）和人工评估结果，全面验证系统效果",
        "完整的课程设计报告：系统阐述大语言模型实现智能问答的技术原理、系统设计方案、实现细节与测试结果",
    ]
    for i, outcome in enumerate(outcomes):
        add_paragraph_with_format(doc, f"（{i+1}）{outcome}", font_size=14, font_name="宋体",
                                  first_line_indent=0.74, space_after=2)
    
    doc.add_page_break()
    
    # ===== 六、参考文献 =====
    add_paragraph_with_format(doc, "六、参考文献", font_size=18, bold=True,
                              font_name="宋体", space_before=14, space_after=10)
    
    refs = [
        "[1] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]// Advances in Neural Information Processing Systems (NeurIPS). 2017, 30: 5998-6008.",
        "[2] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]// Proceedings of NAACL-HLT. 2019: 4171-4186.",
        "[3] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]// Advances in Neural Information Processing Systems (NeurIPS). 2020, 33: 1877-1901.",
        "[4] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]// Advances in Neural Information Processing Systems (NeurIPS). 2020, 33: 9459-9474.",
        "[5] Kamalloo E, Dziri N, Clarke C L A, et al. Evaluating Open-Domain Question Answering in the Era of Large Language Models[C]// Proceedings of ACL. 2023: 5591-5606.",
        "[6] Chen D, Fisch A, Weston J, et al. Reading Wikipedia to Answer Open-Domain Questions[C]// Proceedings of ACL. 2017: 1870-1879.",
        "[7] Izacard G, Grave E. Leveraging Passage Retrieval with Generative Models for Open Domain Question Answering[C]// Proceedings of EACL. 2021: 874-880.",
        "[8] Rajpurkar P, Zhang J, Lopyrev K, et al. SQuAD: 100,000+ Questions for Machine Comprehension of Text[C]// Proceedings of EMNLP. 2016: 2383-2392.",
    ]
    for ref in refs:
        add_paragraph_with_format(doc, ref, font_size=12, font_name="宋体",
                                  first_line_indent=0, space_after=2)


def main():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 创建封面
    create_cover_page(doc)
    
    # 创建目录
    create_toc_page(doc)
    doc.add_page_break()
    
    # 创建正文
    create_body_content(doc)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "中期报告_基于大语言模型的智能问答方法研究与实现.docx")
    doc.save(output_path)
    print(f"[OK] Word document generated: {output_path}")


if __name__ == "__main__":
    main()
